"""Build the colleague-facing Word email for the crop varieties dataset.
Run: python build_email_docx.py
Output: Crop_Varieties_Access_Email.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# --- Base style ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def add_h(text, size=16, bold=True, space_before=12, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p

def add_p(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        p.add_run(text)
    else:
        r.text = text
    return p

def add_numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p

# ============ HEADER ============
add_h("Subject: Crop variety catalogue — ask Claude Code anything, ~10 min setup",
      size=14, space_before=0, space_after=10)

add_p("Hi team,")
add_p(
    "I've pulled together a single unified table of public crop variety "
    "registrations and releases — 45,670 rows, 70 countries, 8 sources, "
    "1935 to 2026. It lives in our Databricks workspace and you can query it "
    "in plain English via Claude Code. No SQL needed."
)
add_p("Below is everything you need to get up and running. If you get stuck at any step, ping me.")

# ============ WHAT YOU GET ============
add_h("What you get")
add_bullet("A single Databricks table: ggo_agdev.agdev.ref_varieties")
add_bullet("Coverage: India 22k, sub-Saharan Africa 11k, the rest of the world 12k")
add_bullet("Major crops: maize, rice, wheat, sorghum, cowpea, cassava, banana, plus 240+ others")
add_bullet("Provenance for every row (source URL + retrieval timestamp)")
add_bullet("Country-level inference flags so you know what to trust")

# ============ SETUP ============
add_h("Setup — three commands, once per machine")

add_p(
    "Both tools are free and small. Open a terminal (PowerShell on Windows, "
    "Terminal on Mac) and run these:",
    italic=True
)

add_p("1) Install Claude Code CLI", bold=True)
add_p("Windows (PowerShell):")
add_code("winget install Anthropic.Claude")
add_p("Mac:")
add_code("brew install --cask claude-code")

add_p("2) Install Databricks CLI", bold=True)
add_p("Windows (PowerShell):")
add_code("winget install Databricks.DatabricksCLI")
add_p("Mac:")
add_code("brew tap databricks/tap && brew install databricks")

add_p("3) Log in to Databricks", bold=True)
add_p("Run this and follow the browser prompt — use your @gatesfoundation.org account:")
add_code("databricks auth login --host https://adb-312788264243632.12.azuredatabricks.net")
add_p(
    "That's it for setup. The login is good for weeks; you only repeat this step "
    "when it eventually expires.",
    italic=True
)

# ============ HOW TO USE ============
add_h("How to use it — just ask")

add_p("In any folder, type:")
add_code("claude")
add_p(
    "You'll get a chat prompt. Ask your question in plain English. Claude has the "
    "Databricks CLI on hand and will write the SQL, run it, and read the results back."
)
add_p("A few things to try:", bold=True)
for q in [
    'Which maize varieties were registered in Kenya since 2023?',
    'How many cassava varieties does Nigeria have in this catalogue?',
    'Compare wheat variety counts across India, Pakistan, and Bangladesh.',
    'Show me Ethiopia rice releases with year, breeder, and release status.',
    'Which CIMMYT maize lines target Eastern Africa? Group by year.',
    'In the last five years, which countries released the most sorghum varieties?',
    'Where do the AGRA catalogue and the national Ghana 2019 catalogue disagree?',
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(q)
    r.italic = True
    p.paragraph_format.space_after = Pt(2)

add_p(
    "Claude will explain what it's about to do before running anything. You can "
    "redirect it at any point — \"actually, restrict that to released-only\" or "
    "\"give me the top 10 instead\" — and it will adjust."
)

# ============ WHERE IT LIVES ============
add_h("Where it lives (for the curious)")
add_bullet("Table: ggo_agdev.agdev.ref_varieties (Databricks, our shared catalog)")
add_bullet("Raw parquet: ggo_agdev.agdev.staging Volume → crop_varieties_canonical/varieties.parquet")
add_bullet("Source code + methodology: GitHub repo (link below)")
add_bullet("Refresh: live API sources auto-refresh quarterly; PDF sources when their publishers re-issue")

add_p("Catalog Explorer link (Databricks UI):")
add_code(
    "https://adb-312788264243632.12.azuredatabricks.net/"
    "explore/data/ggo_agdev/agdev/ref_varieties"
)

# ============ THINGS WORTH KNOWING ============
add_h("Things worth knowing")
add_bullet(
    "Provenance is preserved. Every row has source + source_url + retrieved_at. "
    "No black-box magic."
)
add_bullet(
    "Some country tags are inferred. About 700 rows had missing or placeholder "
    "country fields; I inferred them from the breeder organisation (e.g. 'KALRO' "
    "→ Kenya). Filter with country_inferred = false if you want to exclude these."
)
add_bullet(
    "153 rows have quality flags in review_flags — mostly minor edge cases. "
    "Filter them out if you need a clean cut."
)
add_bullet(
    "PDF sources are point-in-time. KEPHIS reflects Feb 2025, NACGRAB reflects "
    "April 2025, etc. Live API sources (AGRA, India PPV&FRA, FAO WIEWS, CIMMYT) "
    "stay current."
)

# ============ TROUBLESHOOTING ============
add_h("If something goes wrong")
add_bullet(
    "\"databricks: command not found\" — restart your terminal after installing, "
    "so it picks up the new PATH."
)
add_bullet(
    "\"403 Forbidden\" or \"permission denied\" when querying — message me, you "
    "probably need a grant on the ggo_agdev catalog."
)
add_bullet(
    "Claude Code says it can't find the table — make sure you ran "
    "databricks auth login successfully (the command above shows your user when it works)."
)
add_bullet(
    "Anything else — drop me a line. I'd rather help fix the friction than have "
    "you struggle for an afternoon."
)

# ============ FOOTER ============
add_h("Repo + license", size=13)
add_p("GitHub:")
add_code("https://github.com/gatesfoundation/crop_varieties_canonical")
add_p(
    "(Note: the repo currently lives on my personal GitHub account during the "
    "rollout; it'll transfer to gatesfoundation/ shortly. The Databricks table "
    "above is stable regardless.)",
    italic=True
)
add_p("License: CC-BY-4.0 on the unification work; underlying sources keep their original licenses.")

add_p("Happy querying.")
add_p("— Neil")

out = r"C:\Users\neilha\Downloads\crop_varieties_canonical\Crop_Varieties_Access_Email.docx"
doc.save(out)
print(f"Wrote {out}")
